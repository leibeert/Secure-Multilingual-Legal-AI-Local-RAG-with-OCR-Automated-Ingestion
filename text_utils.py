import pdfplumber
import docx
import re
import os
import arabic_reshaper
from bidi.algorithm import get_display
from langchain_core.documents import Document

def fix_arabic_text(text):
    """
    Fixes reversed/disjointed Arabic text (e.g. txeT -> Text).
    """
    if not text: return ""
    # Check if text contains Arabic characters
    if re.search(r'[\u0600-\u06FF]', text):
        try:
            # Reshape (connect letters) then Bidi (fix direction)
            return get_display(arabic_reshaper.reshape(text))
        except:
            return text
    return text

def clean_text(text):
    """
    Removes artifacts like page numbers or repeated headers.
    """
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.isdigit(): continue  # Skip standalone page numbers
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

def split_text_by_articles(full_text, source_name):
    """
    Splits text strictly by 'Article X' or 'المادة X'.
    """
    # Regex captures the delimiter to keep the Title (Article 1) attached
    pattern = r'(^|\n)(Article\s+\d+|المادة\s+[\u0600-\u06FF0-9]+)'
    segments = re.split(pattern, full_text)
    
    docs = []
    current_title = "Preamble/Introduction"
    subject = os.path.splitext(source_name)[0]  # Filename is the Subject

    for segment in segments:
        segment = segment.strip()
        if not segment: continue
        
        # If segment is a Header (Article 5), update title
        if re.match(r'^(Article\s+\d+|المادة\s+[\u0600-\u06FF0-9]+)', segment):
            current_title = segment
        else:
            # It's body text. Create Document.
            if len(segment) > 20: 
                docs.append(Document(
                    page_content=f"{current_title}\n{segment}", # Keep title in content
                    metadata={
                        "source": source_name,
                        "subject": subject,
                        "article": current_title
                    }
                ))
    return docs

def load_file_structured(file_path):
    """
    Main loader function.
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    full_text = ""

    # 1. Extract Full Text
    if file_ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                raw = page.extract_text()
                fixed = fix_arabic_text(raw)
                cleaned = clean_text(fixed)
                full_text += cleaned + "\n"
                
    elif file_ext == ".docx":
        doc = docx.Document(file_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        # Note: DOCX usually doesn't need Arabic reshaping, but test it.
    
    # 2. Split by Articles
    return split_text_by_articles(full_text, filename)