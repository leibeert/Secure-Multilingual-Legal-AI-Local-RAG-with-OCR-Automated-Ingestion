import os
import shutil
import re
import pickle
import fitz  # PyMuPDF
import docx
from langchain_chroma import Chroma
from langchain.storage import InMemoryStore
from langchain.retrievers import ParentDocumentRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_core.documents import Document

# --- CONFIGURATION ---
DATA_FOLDER = "data"
OUTPUT_DIR = "saudi_legal_db_final1"

# Setup Folders
# Setup Folders moved to main execution to allow safe importing


# --- ROBUST TEXT EXTRACTION ---

def clean_text(text):
    """
    Cleans up Arabic text artifacts.
    """
    # Fix common Arabic encoding issues
    text = text.replace('\xa0', ' ')  # Non-breaking space
    
    lines = text.split('\n')
    cleaned = []
    for line in lines:
        line = line.strip()
        if not line: continue
        # Skip simple page numbers (1-3 digits)
        if line.isdigit() and len(line) < 4: continue
        cleaned.append(line)
    return '\n'.join(cleaned)

def smart_split(text, source):
    """
    Matches:
    1. Article 1 (Digits)
    2. المادة 1 (Digits)
    3. المادة الأولى (Words)
    Include Presentation Forms: 'اﻟﻤﺎدة'
    """
    # Pattern explanation:
    # (?:Article|المادة|اﻟﻤﺎدة) : Match various forms of "Article"
    # \s+ : Space
    # (?: ... ) : Match Number/Word
    # \d+ : Digits
    # | : OR
    # \w+ : Any word char (covers standard Arabic AND Presentation forms)
    pattern = r'(?:^|\n)((?:Article|المادة|اﻟﻤﺎدة)\s+(?:\d+|\w+(?:\s+\w+){0,4}))'
    
    segments = re.split(pattern, text)
    docs = []
    current_title = "Introduction"
    subject = os.path.splitext(source)[0]
    
    found_articles = 0

    for seg in segments:
        seg = seg.strip()
        if not seg: continue
        
        # Check if header
        if re.match(r'^(Article|المادة|اﻟﻤﺎدة)', seg):
            current_title = seg
            found_articles += 1
        # Check if body
        elif len(seg) > 20:
            docs.append(Document(
                page_content=f"Source: {source}\nSection: {current_title}\n\n{seg}",
                metadata={"source": source, "subject": subject, "article": current_title}
            ))
            
    # Fallback: If we found text but no "Articles", chunk by size
    if found_articles < 2 and len(text) > 1000:
        print(f"      ⚠️ Regex found only {found_articles} articles. Switching to Standard Chunking.")
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        raw_docs = splitter.create_documents([text])
        docs = []
        for i, d in enumerate(raw_docs):
            d.metadata = {"source": source, "subject": subject, "article": f"Page/Part {i+1}"}
            d.page_content = f"Source: {source}\nPart: {i+1}\n\n{d.page_content}"
            docs.append(d)
        
    return docs

def load_file(filepath):
    name = os.path.basename(filepath)
    ext = os.path.splitext(filepath)[1].lower()
    full_text = ""
    
    try:
        if ext == '.pdf':
            # USE PYMUPDF (FITZ) - More robust than pdfplumber
            doc = fitz.open(filepath)
            
            # DEBUG: Print first 100 chars to see if it works
            if len(doc) > 0:
                first_page = doc[0].get_text()
                print(f"      👀 FITZ Sees: {first_page[:100].replace(chr(10), ' ')}...")
            
            for page in doc:
                # "text" mode is standard, "blocks" can be better for columns
                full_text += page.get_text() + "\n"
                
        elif ext == '.docx':
            doc = docx.Document(filepath)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
        full_text = clean_text(full_text)
        print(f"   📖 Extracted {len(full_text)} chars from {name}...")
        
        return smart_split(full_text, name)
        
    except Exception as e:
        print(f"   ❌ Error reading {name}: {e}")
        return []

# --- MAIN EXECUTION ---
def main():
    import torch
    device = "cuda" if torch.cuda.is_available() else "cpu"
    # Setup Folders (Destructive - Only run when script is executed directly)
    if os.path.exists(OUTPUT_DIR):
        shutil.rmtree(OUTPUT_DIR)
    os.makedirs(OUTPUT_DIR)
    os.makedirs(DATA_FOLDER, exist_ok=True)

    print(f"🔌 Loading BAAI/bge-m3 on {device.upper()}...")
    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-m3",
        model_kwargs={'device': device}, 
        encode_kwargs={'normalize_embeddings': True}
    )

    print("🔨 Creating ChromaDB...")
    vector_db_path = os.path.join(OUTPUT_DIR, "chroma_vectors")
    vectorstore = Chroma(
        collection_name="saudi_legal_docs",
        embedding_function=embeddings,
        persist_directory=vector_db_path
    )
    docstore = InMemoryStore()

    retriever = ParentDocumentRetriever(
        vectorstore=vectorstore,
        docstore=docstore,
        child_splitter=RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=100),
    )

    print("\n🚀 Starting PyMuPDF Ingestion (V6)...")
    files = [f for f in os.listdir(DATA_FOLDER) if f.endswith(('.pdf', '.docx'))]
    
    if not files:
        print(f"⚠️ No files found in '{DATA_FOLDER}'!")
        return

    total_chunks = 0
    for f in files:
        path = os.path.join(DATA_FOLDER, f)
        print(f"   📄 Processing: {f}...", end="\n")
        
        docs = load_file(path)
        if docs:
            retriever.add_documents(docs, ids=None)
            count = len(docs)
            total_chunks += count
            print(f"      ✅ Indexed {count} sections.")
        else:
            print("      ⚠️ Skipped (Empty).")

    print(f"\n💾 Saving Document Store (Total {total_chunks} items)...")
    with open(os.path.join(OUTPUT_DIR, "docstore.pkl"), "wb") as f:
        pickle.dump(docstore, f)

   
    print("🎉 DONE! .")

if __name__ == "__main__":
    main()